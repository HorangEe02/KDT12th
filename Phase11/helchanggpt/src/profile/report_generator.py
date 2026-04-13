"""
report_generator.py
인바디 분석 결과를 다양한 파일 포맷으로 내보냅니다.

지원 포맷: CSV, PDF, Word(.docx), Excel(.xlsx), JSON
"""

import json
from datetime import datetime
from io import BytesIO
from pathlib import Path

from .inbody_parser import InBodyData
from .inbody_analyzer import InBodyAnalysis


# ═══════════════════════════════════════
# CSV 내보내기
# ═══════════════════════════════════════

def export_to_csv(
    inbody: InBodyData,
    output_path: str | None = None,
) -> str | BytesIO:
    """인바디 데이터를 CSV로 내보냅니다."""
    import pandas as pd

    row = {
        "측정일": inbody.basic_info.get("test_date", ""),
        "키(cm)": inbody.basic_info.get("height_cm"),
        "나이": inbody.basic_info.get("age"),
        "성별": inbody.basic_info.get("gender"),
        "체중(kg)": inbody.body_composition.get("weight_kg"),
        "골격근량(kg)": inbody.muscle_fat_analysis.get("skeletal_muscle_mass_kg"),
        "체지방량(kg)": inbody.muscle_fat_analysis.get("body_fat_mass_kg")
            or inbody.body_composition.get("body_fat_mass_kg"),
        "체지방률(%)": inbody.obesity_analysis.get("body_fat_percent"),
        "BMI": inbody.obesity_analysis.get("bmi"),
        "기초대사량(kcal)": inbody.research_params.get("bmr_kcal"),
        "인바디점수": inbody.inbody_score,
        "내장지방(cm²)": (inbody.visceral_fat or {}).get("vfa_cm2"),
        "SMI(kg/m²)": inbody.research_params.get("smi"),
        "복부지방률": inbody.research_params.get("waist_hip_ratio"),
        "적정체중(kg)": inbody.weight_control.get("ideal_weight_kg"),
        "체중조절(kg)": inbody.weight_control.get("weight_adjustment_kg"),
        "지방조절(kg)": inbody.weight_control.get("fat_adjustment_kg"),
        "근육조절(kg)": inbody.weight_control.get("muscle_adjustment_kg"),
    }

    df = pd.DataFrame([row])

    if output_path:
        df.to_csv(output_path, index=False, encoding="utf-8-sig")
        return output_path

    buffer = BytesIO()
    df.to_csv(buffer, index=False, encoding="utf-8-sig")
    buffer.seek(0)
    return buffer


def export_history_to_csv(
    history_records: list[InBodyData],
    output_path: str | None = None,
) -> str | BytesIO:
    """여러 인바디 측정 기록을 시계열 CSV로 내보냅니다."""
    import pandas as pd

    rows = []
    for record in history_records:
        rows.append({
            "측정일": record.basic_info.get("test_date", ""),
            "체중(kg)": record.body_composition.get("weight_kg"),
            "골격근량(kg)": record.muscle_fat_analysis.get("skeletal_muscle_mass_kg"),
            "체지방량(kg)": record.body_composition.get("body_fat_mass_kg"),
            "체지방률(%)": record.obesity_analysis.get("body_fat_percent"),
            "BMI": record.obesity_analysis.get("bmi"),
            "인바디점수": record.inbody_score,
        })

    df = pd.DataFrame(rows)

    if output_path:
        df.to_csv(output_path, index=False, encoding="utf-8-sig")
        return output_path

    buffer = BytesIO()
    df.to_csv(buffer, index=False, encoding="utf-8-sig")
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════
# Excel 내보내기
# ═══════════════════════════════════════

def export_to_excel(
    inbody: InBodyData,
    analysis: InBodyAnalysis | None = None,
    output_path: str | None = None,
) -> str | BytesIO:
    """인바디 데이터 + 분석 결과를 Excel로 내보냅니다."""
    import pandas as pd

    buffer = BytesIO()

    with pd.ExcelWriter(buffer if not output_path else output_path, engine="openpyxl") as writer:
        # Sheet 1: 측정 데이터
        data_row = {
            "항목": ["체중", "골격근량", "체지방량", "체지방률", "BMI",
                     "기초대사량", "인바디점수", "내장지방", "SMI", "복부지방률"],
            "수치": [
                inbody.body_composition.get("weight_kg"),
                inbody.muscle_fat_analysis.get("skeletal_muscle_mass_kg"),
                inbody.body_composition.get("body_fat_mass_kg"),
                inbody.obesity_analysis.get("body_fat_percent"),
                inbody.obesity_analysis.get("bmi"),
                inbody.research_params.get("bmr_kcal"),
                inbody.inbody_score,
                (inbody.visceral_fat or {}).get("vfa_cm2"),
                inbody.research_params.get("smi"),
                inbody.research_params.get("waist_hip_ratio"),
            ],
            "단위": ["kg", "kg", "kg", "%", "kg/m²", "kcal", "점", "cm²", "kg/m²", ""],
        }
        df_data = pd.DataFrame(data_row)
        df_data.to_excel(writer, sheet_name="측정데이터", index=False)

        # Sheet 2: 분석 결과
        if analysis:
            analysis_rows = {
                "항목": ["체형", "BMI 상태", "체지방률 상태", "종합 판정",
                         "내장지방 위험", "근감소증 위험", "수분 균형", "종합 등급"],
                "결과": [
                    analysis.body_shape, analysis.bmi_status,
                    analysis.body_fat_status, analysis.combined_assessment,
                    analysis.visceral_fat_risk, analysis.sarcopenia_risk,
                    analysis.ecw_status, f"{analysis.overall_grade} ({analysis.inbody_score}점)",
                ],
            }
            df_analysis = pd.DataFrame(analysis_rows)
            df_analysis.to_excel(writer, sheet_name="분석결과", index=False)

        # Sheet 3: 코칭 메시지
        if analysis and analysis.coaching_messages:
            coaching_rows = {
                "구분": (
                    ["칭찬"] * len(analysis.praise_points) +
                    ["개선"] * len(analysis.improvement_points) +
                    ["행동"] * len(analysis.action_items)
                ),
                "내용": analysis.praise_points + analysis.improvement_points + analysis.action_items,
            }
            df_coaching = pd.DataFrame(coaching_rows)
            df_coaching.to_excel(writer, sheet_name="코칭메시지", index=False)

        # Sheet 4: 부위별 근육 (있을 경우)
        if inbody.segmental_lean:
            seg = inbody.segmental_lean
            seg_rows = {
                "부위": ["오른팔", "왼팔", "몸통", "오른다리", "왼다리"],
                "근육량(kg)": [
                    seg.get("right_arm_kg"), seg.get("left_arm_kg"),
                    seg.get("trunk_kg"), seg.get("right_leg_kg"), seg.get("left_leg_kg"),
                ],
                "발달률(%)": [
                    seg.get("right_arm_pct"), seg.get("left_arm_pct"),
                    seg.get("trunk_pct"), seg.get("right_leg_pct"), seg.get("left_leg_pct"),
                ],
            }
            df_seg = pd.DataFrame(seg_rows)
            df_seg.to_excel(writer, sheet_name="부위별근육", index=False)

    if output_path:
        return output_path

    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════
# PDF 보고서
# ═══════════════════════════════════════

def export_to_pdf(
    inbody: InBodyData,
    analysis: InBodyAnalysis | None = None,
    output_path: str | None = None,
) -> str | BytesIO:
    """인바디 분석 보고서를 PDF로 생성합니다."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        # reportlab이 없으면 간단한 텍스트 PDF로 폴백
        return _export_to_pdf_simple(inbody, analysis, output_path)

    buffer = BytesIO()
    target = output_path or buffer

    doc = SimpleDocTemplate(
        target,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )

    styles = getSampleStyleSheet()
    elements = []

    # 한글 폰트 등록 시도
    _register_korean_font()

    # 제목
    elements.append(Paragraph("HelChangGPT 건강 분석 보고서", styles["Title"]))
    elements.append(Spacer(1, 5 * mm))

    # 기본 정보
    bi = inbody.basic_info
    info_text = f"이름: - | 나이: {bi.get('age', '-')}세 | 성별: {bi.get('gender', '-')} | 키: {bi.get('height_cm', '-')}cm | 측정일: {bi.get('test_date', '-')}"
    elements.append(Paragraph(info_text, styles["Normal"]))
    elements.append(Spacer(1, 5 * mm))

    # 핵심 지표 테이블
    bc = inbody.body_composition
    oa = inbody.obesity_analysis
    rp = inbody.research_params
    vf = inbody.visceral_fat or {}

    key_data = [
        ["항목", "수치", "판정"],
        ["체중", f"{bc.get('weight_kg', '-')} kg", ""],
        ["골격근량", f"{inbody.muscle_fat_analysis.get('skeletal_muscle_mass_kg', '-')} kg", ""],
        ["체지방량", f"{bc.get('body_fat_mass_kg', '-')} kg", ""],
        ["체지방률", f"{oa.get('body_fat_percent', '-')} %", analysis.body_fat_status if analysis else ""],
        ["BMI", f"{oa.get('bmi', '-')}", analysis.bmi_status if analysis else ""],
        ["기초대사량", f"{rp.get('bmr_kcal', '-')} kcal", ""],
        ["인바디점수", f"{inbody.inbody_score or '-'} / 100", analysis.overall_grade if analysis else ""],
        ["내장지방", f"{vf.get('vfa_cm2', '-')} cm2", analysis.visceral_fat_risk if analysis else ""],
        ["SMI", f"{rp.get('smi', '-')} kg/m2", analysis.sarcopenia_risk if analysis else ""],
    ]

    table = Table(key_data, colWidths=[50 * mm, 40 * mm, 60 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 8 * mm))

    # 분석 결과
    if analysis:
        elements.append(Paragraph("종합 분석", styles["Heading2"]))
        elements.append(Paragraph(f"체형: {analysis.body_shape} - {analysis.body_shape_description}", styles["Normal"]))
        elements.append(Paragraph(f"종합: {analysis.combined_assessment}", styles["Normal"]))
        elements.append(Spacer(1, 5 * mm))

        # 코칭 메시지
        if analysis.praise_points:
            elements.append(Paragraph("잘하고 있는 점:", styles["Heading3"]))
            for p in analysis.praise_points:
                elements.append(Paragraph(f"  + {p}", styles["Normal"]))

        if analysis.improvement_points:
            elements.append(Paragraph("개선 필요:", styles["Heading3"]))
            for p in analysis.improvement_points:
                elements.append(Paragraph(f"  - {p}", styles["Normal"]))

        if analysis.action_items:
            elements.append(Paragraph("추천 행동:", styles["Heading3"]))
            for a in analysis.action_items:
                elements.append(Paragraph(f"  > {a}", styles["Normal"]))

    # 푸터
    elements.append(Spacer(1, 10 * mm))
    elements.append(Paragraph(
        f"생성: HelChangGPT v1.0 | {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        styles["Normal"],
    ))

    doc.build(elements)

    if output_path:
        return output_path

    buffer.seek(0)
    return buffer


def _export_to_pdf_simple(inbody, analysis, output_path) -> str | BytesIO:
    """reportlab 없이 간단한 텍스트 기반 PDF 생성 (FPDF 사용)."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError("PDF 생성에 reportlab 또는 fpdf2 패키지가 필요합니다.\npip install reportlab 또는 pip install fpdf2")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)

    pdf.cell(0, 10, "HelChangGPT Health Analysis Report", ln=True, align="C")
    pdf.ln(5)

    bi = inbody.basic_info
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 7, f"Age: {bi.get('age', '-')} | Gender: {bi.get('gender', '-')} | Height: {bi.get('height_cm', '-')}cm", ln=True)
    pdf.ln(3)

    # 핵심 지표
    metrics = [
        ("Weight", f"{inbody.body_composition.get('weight_kg', '-')} kg"),
        ("SMM", f"{inbody.muscle_fat_analysis.get('skeletal_muscle_mass_kg', '-')} kg"),
        ("BFM", f"{inbody.body_composition.get('body_fat_mass_kg', '-')} kg"),
        ("BF%", f"{inbody.obesity_analysis.get('body_fat_percent', '-')} %"),
        ("BMI", f"{inbody.obesity_analysis.get('bmi', '-')}"),
        ("BMR", f"{inbody.research_params.get('bmr_kcal', '-')} kcal"),
        ("Score", f"{inbody.inbody_score or '-'} / 100"),
    ]

    for name, value in metrics:
        pdf.cell(50, 7, name, border=1)
        pdf.cell(50, 7, value, border=1, ln=True)

    if analysis:
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, "Analysis", ln=True)
        pdf.set_font("Helvetica", size=9)
        pdf.cell(0, 6, f"Body Shape: {analysis.body_shape}", ln=True)
        pdf.cell(0, 6, f"Assessment: {analysis.combined_assessment}", ln=True)

    buffer = BytesIO()
    pdf_bytes = pdf.output()
    buffer.write(pdf_bytes)
    buffer.seek(0)

    if output_path:
        with open(output_path, "wb") as f:
            f.write(buffer.read())
        return output_path

    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════
# Word 보고서
# ═══════════════════════════════════════

def export_to_docx(
    inbody: InBodyData,
    analysis: InBodyAnalysis | None = None,
    output_path: str | None = None,
) -> str | BytesIO:
    """인바디 분석 보고서를 Word(.docx)로 생성합니다."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()

    # 제목
    title = doc.add_heading("HelChangGPT 건강 분석 보고서", level=0)

    # 기본 정보
    bi = inbody.basic_info
    doc.add_paragraph(
        f"나이: {bi.get('age', '-')}세 | 성별: {bi.get('gender', '-')} | "
        f"키: {bi.get('height_cm', '-')}cm | 측정일: {bi.get('test_date', '-')}"
    )

    # 핵심 지표 테이블
    doc.add_heading("핵심 지표", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"

    headers = table.rows[0].cells
    headers[0].text = "항목"
    headers[1].text = "수치"
    headers[2].text = "판정"

    bc = inbody.body_composition
    oa = inbody.obesity_analysis
    rp = inbody.research_params
    vf = inbody.visceral_fat or {}

    metrics = [
        ("체중", f"{bc.get('weight_kg', '-')} kg", ""),
        ("골격근량", f"{inbody.muscle_fat_analysis.get('skeletal_muscle_mass_kg', '-')} kg", ""),
        ("체지방량", f"{bc.get('body_fat_mass_kg', '-')} kg", ""),
        ("체지방률", f"{oa.get('body_fat_percent', '-')} %", analysis.body_fat_status if analysis else ""),
        ("BMI", f"{oa.get('bmi', '-')}", analysis.bmi_status if analysis else ""),
        ("기초대사량", f"{rp.get('bmr_kcal', '-')} kcal", ""),
        ("인바디점수", f"{inbody.inbody_score or '-'} / 100", analysis.overall_grade if analysis else ""),
        ("내장지방", f"{vf.get('vfa_cm2', '-')} cm²", analysis.visceral_fat_risk if analysis else ""),
        ("SMI", f"{rp.get('smi', '-')} kg/m²", analysis.sarcopenia_risk if analysis else ""),
    ]

    for name, value, status in metrics:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = value
        row[2].text = status

    # 분석 결과
    if analysis:
        doc.add_heading("종합 분석", level=1)
        doc.add_paragraph(f"체형: {analysis.body_shape} — {analysis.body_shape_description}")
        doc.add_paragraph(f"종합 판정: {analysis.combined_assessment}")

        if analysis.praise_points:
            doc.add_heading("잘하고 있는 점", level=2)
            for p in analysis.praise_points:
                doc.add_paragraph(p, style="List Bullet")

        if analysis.improvement_points:
            doc.add_heading("개선이 필요한 점", level=2)
            for p in analysis.improvement_points:
                doc.add_paragraph(p, style="List Bullet")

        if analysis.action_items:
            doc.add_heading("추천 행동", level=2)
            for a in analysis.action_items:
                doc.add_paragraph(a, style="List Number")

    # 푸터
    doc.add_paragraph("")
    footer = doc.add_paragraph(
        f"생성: HelChangGPT v1.0 | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    if output_path:
        doc.save(output_path)
        return output_path

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════
# JSON 내보내기
# ═══════════════════════════════════════

def export_to_json(
    inbody: InBodyData,
    analysis: InBodyAnalysis | None = None,
    output_path: str | None = None,
) -> str | dict:
    """인바디 데이터 + 분석 결과를 JSON으로 내보냅니다."""
    data = {
        "inbody_data": inbody.raw,
        "parse_method": inbody.parse_method,
        "parse_confidence": inbody.parse_confidence,
        "validation": inbody.validation,
        "exported_at": datetime.now().isoformat(),
    }

    if analysis:
        data["analysis"] = {
            "body_shape": analysis.body_shape,
            "body_shape_description": analysis.body_shape_description,
            "bmi_status": analysis.bmi_status,
            "body_fat_status": analysis.body_fat_status,
            "combined_assessment": analysis.combined_assessment,
            "visceral_fat_risk": analysis.visceral_fat_risk,
            "sarcopenia_risk": analysis.sarcopenia_risk,
            "ecw_status": analysis.ecw_status,
            "overall_grade": analysis.overall_grade,
            "inbody_score": analysis.inbody_score,
            "praise_points": analysis.praise_points,
            "improvement_points": analysis.improvement_points,
            "action_items": analysis.action_items,
            "coaching_messages": analysis.coaching_messages,
        }

    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return output_path

    return data


# ═══════════════════════════════════════
# 통합 내보내기
# ═══════════════════════════════════════

def export_report(
    inbody: InBodyData,
    analysis: InBodyAnalysis | None = None,
    format: str = "pdf",
    output_path: str | None = None,
) -> str | BytesIO | dict:
    """
    지정된 포맷으로 보고서를 생성합니다.

    Args:
        format: "csv" | "pdf" | "docx" | "xlsx" | "json"
    """
    exporters = {
        "csv": export_to_csv,
        "pdf": export_to_pdf,
        "docx": export_to_docx,
        "xlsx": export_to_excel,
        "json": export_to_json,
    }

    exporter = exporters.get(format)
    if not exporter:
        raise ValueError(f"지원하지 않는 포맷: {format} (지원: {list(exporters.keys())})")

    if format in ("xlsx", "pdf", "docx", "json") and analysis:
        return exporter(inbody, analysis, output_path)
    else:
        return exporter(inbody, output_path=output_path)


def _register_korean_font():
    """한글 PDF 출력을 위한 폰트 등록 시도."""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import platform

        if platform.system() == "Darwin":
            font_paths = [
                "/System/Library/Fonts/AppleSDGothicNeo.ttc",
                "/Library/Fonts/NanumGothic.ttf",
            ]
        else:
            font_paths = [
                "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            ]

        for fp in font_paths:
            if Path(fp).exists():
                pdfmetrics.registerFont(TTFont("Korean", fp))
                return
    except Exception:
        pass
