"""
inbody_parser.py
인바디(InBody) 체성분 분석 결과지에서 데이터를 추출합니다.

지원 모델: InBody 770, InBody 970 등
지원 입력: 이미지(PNG/JPG), PDF, CSV, Word(.docx), Excel(.xlsx)

방식 A: LLM Vision (Ollama gemma4 / OpenAI GPT-4o) — 이미지 파싱
방식 B: 파일 파싱 (CSV/Excel/PDF/Word) — 구조화 데이터 직접 읽기
방식 C: OCR + 정규식 — 보조/비교용
"""

import base64
import json
import math
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path


@dataclass
class InBodyData:
    """파싱된 인바디 데이터 (InBody 970 전체 필드 지원)"""
    raw: dict
    basic_info: dict                         # 기본 정보
    body_composition: dict                   # 체성분 분석
    muscle_fat_analysis: dict                # 골격근·지방 분석
    obesity_analysis: dict                   # 비만 분석
    weight_control: dict                     # 체중 조절
    research_params: dict                    # 연구항목
    segmental_lean: dict = field(default_factory=dict)   # 부위별 근육분석
    ecw_ratio: dict = field(default_factory=dict)        # 세포외수분비
    body_history: dict = field(default_factory=dict)     # 신체변화 히스토리
    visceral_fat: dict = field(default_factory=dict)     # 내장지방단면적
    inbody_score: int | None = None
    parse_method: str = ""
    parse_confidence: float = 0.0
    validation: dict = field(default_factory=dict)       # 교차 검증 결과


# ═══════════════════════════════════════
# InBody 970 전체 필드 파싱 프롬프트
# ═══════════════════════════════════════

INBODY_PARSE_PROMPT_V2 = """당신은 InBody(인바디) 체성분 분석 결과지를 정확하게 읽는 전문가입니다.
이미지에서 모든 수치를 정확히 추출하여 아래 JSON 형식으로 반환하세요.
숫자는 이미지에 표시된 그대로 입력하세요. 없는 항목은 null로 표시하세요.

{
  "basic_info": {
    "height_cm": <키(cm)>,
    "age": <나이>,
    "gender": "<남성 또는 여성>",
    "test_date": "<검사일시 YYYY-MM-DD HH:MM>",
    "model": "<InBody 모델명 (예: InBody970, InBody770)>"
  },
  "body_composition": {
    "body_water_L": <체수분(L)>,
    "intracellular_water_L": <세포내수분(L)>,
    "extracellular_water_L": <세포외수분(L)>,
    "protein_kg": <단백질(kg)>,
    "minerals_kg": <무기질(kg)>,
    "body_fat_mass_kg": <체지방(kg)>,
    "weight_kg": <체중(kg)>,
    "lean_body_mass_kg": <근육량(kg)>
  },
  "muscle_fat_analysis": {
    "weight_kg": <체중(kg)>,
    "skeletal_muscle_mass_kg": <골격근량(kg)>,
    "body_fat_mass_kg": <체지방량(kg)>
  },
  "obesity_analysis": {
    "bmi": <BMI>,
    "body_fat_percent": <체지방률(%)>
  },
  "weight_control": {
    "ideal_weight_kg": <적정체중(kg)>,
    "weight_adjustment_kg": <체중조절(kg)>,
    "fat_adjustment_kg": <지방조절(kg)>,
    "muscle_adjustment_kg": <근육조절(kg)>
  },
  "research_params": {
    "intracellular_water_L": <세포내수분(L)>,
    "extracellular_water_L": <세포외수분(L)>,
    "bmr_kcal": <기초대사량(kcal)>,
    "waist_hip_ratio": <복부지방률>,
    "body_cell_mass_kg": <체세포량(kg)>,
    "smi": <SMI(kg/m²)>
  },
  "visceral_fat": {
    "vfa_cm2": <내장지방단면적(cm²)>
  },
  "segmental_lean": {
    "right_arm_kg": <오른팔 근육량(kg)>,
    "left_arm_kg": <왼팔 근육량(kg)>,
    "trunk_kg": <몸통 근육량(kg)>,
    "right_leg_kg": <오른다리 근육량(kg)>,
    "left_leg_kg": <왼다리 근육량(kg)>,
    "right_arm_pct": <오른팔 발달률(%)>,
    "left_arm_pct": <왼팔 발달률(%)>,
    "trunk_pct": <몸통 발달률(%)>,
    "right_leg_pct": <오른다리 발달률(%)>,
    "left_leg_pct": <왼다리 발달률(%)>
  },
  "ecw_ratio": {
    "whole_body": <전신 세포외수분비>,
    "right_arm": <오른팔>,
    "left_arm": <왼팔>,
    "trunk": <몸통>,
    "right_leg": <오른다리>,
    "left_leg": <왼다리>
  },
  "body_history": {
    "dates": ["<날짜1>", "<날짜2>", ...],
    "weight_series": [<체중1>, <체중2>, ...],
    "smm_series": [<골격근량1>, <골격근량2>, ...],
    "bfp_series": [<체지방률1>, <체지방률2>, ...],
    "ecw_series": [<세포외수분비1>, <세포외수분비2>, ...]
  },
  "inbody_score": <인바디점수>
}

JSON만 반환하세요. 다른 설명은 넣지 마세요."""


# ═══════════════════════════════════════
# 이미지 유틸리티
# ═══════════════════════════════════════

def encode_image_to_base64(image_path: str) -> str:
    """이미지를 base64로 인코딩합니다."""
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def _clean_json_response(text: str) -> dict:
    """LLM 응답에서 JSON을 추출합니다."""
    text = text.strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[1] if "\n" in text else text[3:]
        text = text.rsplit("```", 1)[0]
    return json.loads(text)


# ═══════════════════════════════════════
# 교차 검증
# ═══════════════════════════════════════

def validate_inbody_data(data: dict) -> dict:
    """
    파싱된 인바디 데이터의 내부 일관성을 교차 검증합니다.

    검증 항목:
    1. 체중 = 체수분 + 단백질 + 무기질 + 체지방
    2. BMI = 체중 / 키(m)²
    3. 체지방률 = 체지방량 / 체중 × 100
    4. 골격근량 < 체중
    """
    checks = {}
    bc = data.get("body_composition", {})
    oa = data.get("obesity_analysis", {})
    bi = data.get("basic_info", {})
    mfa = data.get("muscle_fat_analysis", {})

    # 1. 체중 구성 검증 (모든 값을 float로 안전 변환)
    def _sf(val):
        if val is None: return 0
        try: return float(val)
        except (ValueError, TypeError): return 0

    water = _sf(bc.get("body_water_L"))
    protein = _sf(bc.get("protein_kg"))
    minerals = _sf(bc.get("minerals_kg"))
    fat = _sf(bc.get("body_fat_mass_kg"))
    weight = _sf(bc.get("weight_kg")) or _sf(mfa.get("weight_kg"))

    if water and protein and minerals and fat and weight:
        calc_weight = water + protein + minerals + fat
        diff = abs(calc_weight - weight)
        checks["weight_composition"] = {
            "expected": weight,
            "calculated": round(calc_weight, 1),
            "difference": round(diff, 1),
            "pass": diff < 1.0,
        }

    # 2. BMI 검증
    height = _safe_float(bi.get("height_cm")) if hasattr(_safe_float, '__call__') and 'bi' in dir() else None
    try:
        height = float(bi.get("height_cm")) if bi.get("height_cm") is not None else None
    except (ValueError, TypeError):
        height = None
    try:
        bmi = float(oa.get("bmi")) if oa.get("bmi") is not None else None
    except (ValueError, TypeError):
        bmi = None
    if height and weight and bmi:
        calc_bmi = round(weight / ((height / 100) ** 2), 1)
        diff = abs(calc_bmi - bmi)
        checks["bmi"] = {
            "expected": bmi,
            "calculated": calc_bmi,
            "difference": round(diff, 1),
            "pass": diff < 0.5,
        }

    # 3. 체지방률 검증
    try:
        bf_pct = float(oa.get("body_fat_percent")) if oa.get("body_fat_percent") is not None else None
    except (ValueError, TypeError):
        bf_pct = None
    if fat and weight and bf_pct:
        calc_pct = round(fat / weight * 100, 1)
        diff = abs(calc_pct - bf_pct)
        checks["body_fat_percent"] = {
            "expected": bf_pct,
            "calculated": calc_pct,
            "difference": round(diff, 1),
            "pass": diff < 1.0,
        }

    # 4. 골격근량 범위
    try:
        smm = float(mfa.get("skeletal_muscle_mass_kg")) if mfa.get("skeletal_muscle_mass_kg") is not None else None
    except (ValueError, TypeError):
        smm = None
    if smm and weight:
        ratio = smm / weight
        checks["smm_ratio"] = {
            "ratio": round(ratio, 3),
            "pass": 0.2 < ratio < 0.7,
        }

    all_pass = all(c.get("pass", True) for c in checks.values())
    checks["all_pass"] = all_pass

    return checks


# ═══════════════════════════════════════
# 방식 A: LLM Vision 파싱
# ═══════════════════════════════════════

def parse_inbody_with_ollama(
    image_path: str,
    model: str = "gemma4:e4b",
    temperature: float = 0.1,
    ollama_base_url: str = "http://localhost:11434/v1",
) -> InBodyData:
    """
    Ollama 비전 모델로 인바디 이미지를 파싱합니다.

    Args:
        image_path: 인바디 이미지 파일 경로
        model: Ollama 비전 모델 (gemma4, gemma4:26b 등)
        temperature: 낮을수록 정확한 수치 추출
        ollama_base_url: Ollama API URL
    """
    from openai import OpenAI

    client = OpenAI(base_url=ollama_base_url, api_key="ollama")
    base64_image = encode_image_to_base64(image_path)

    response = client.chat.completions.create(
        model=model,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": INBODY_PARSE_PROMPT_V2},
                {"type": "image_url", "image_url": {
                    "url": f"data:image/png;base64,{base64_image}",
                    "detail": "high",
                }},
            ],
        }],
        temperature=temperature,
        max_tokens=3000,
    )

    result = _clean_json_response(response.choices[0].message.content)
    return _build_inbody_data(result, f"ollama_{model}")


def parse_inbody_with_openai(
    image_path: str,
    api_key: str,
    model: str = "gpt-4o",
    temperature: float = 0.1,
) -> InBodyData:
    """OpenAI GPT-4o Vision으로 인바디 이미지를 파싱합니다."""
    from openai import OpenAI

    client = OpenAI(api_key=api_key)
    base64_image = encode_image_to_base64(image_path)

    response = client.chat.completions.create(
        model=model,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": INBODY_PARSE_PROMPT_V2},
                {"type": "image_url", "image_url": {
                    "url": f"data:image/png;base64,{base64_image}",
                    "detail": "high",
                }},
            ],
        }],
        temperature=temperature,
        max_tokens=3000,
    )

    result = _clean_json_response(response.choices[0].message.content)
    return _build_inbody_data(result, f"openai_{model}")


# ═══════════════════════════════════════
# 방식 B: 파일 파싱 (CSV/Excel/PDF/Word)
# ═══════════════════════════════════════

def parse_inbody_from_csv(file_path: str) -> InBodyData:
    """CSV 파일에서 인바디 데이터를 읽습니다."""
    import pandas as pd

    df = pd.read_csv(file_path, encoding="utf-8-sig")
    # 컬럼명 정규화
    df.columns = [c.strip() for c in df.columns]

    records = []
    for _, row in df.iterrows():
        record = _map_csv_row_to_inbody(row.to_dict())
        records.append(record)

    if len(records) == 1:
        return _build_inbody_data(records[0], "csv")

    # 여러 행이면 가장 최근 데이터 사용 + 히스토리 구성
    latest = records[-1]
    latest["body_history"] = _build_history_from_records(records)
    return _build_inbody_data(latest, "csv")


def parse_inbody_from_excel(file_path: str) -> InBodyData:
    """Excel 파일에서 인바디 데이터를 읽습니다."""
    import pandas as pd

    df = pd.read_excel(file_path)
    df.columns = [c.strip() for c in df.columns]

    records = []
    for _, row in df.iterrows():
        records.append(_map_csv_row_to_inbody(row.to_dict()))

    latest = records[-1] if records else {}
    if len(records) > 1:
        latest["body_history"] = _build_history_from_records(records)
    return _build_inbody_data(latest, "excel")


def parse_inbody_from_pdf(file_path: str) -> InBodyData:
    """PDF 파일에서 인바디 데이터를 추출합니다."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except ImportError:
            raise ImportError("PyPDF2 또는 pdfplumber가 필요합니다.")

    return _parse_inbody_from_text(text, "pdf")


def parse_inbody_from_docx(file_path: str) -> InBodyData:
    """Word(.docx) 파일에서 인바디 데이터를 추출합니다."""
    from docx import Document

    doc = Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    text = "\n".join(text_parts)
    return _parse_inbody_from_text(text, "docx")


def parse_inbody_auto(file_path: str, **kwargs) -> InBodyData:
    """파일 확장자에 따라 적절한 파서를 자동 선택합니다."""
    ext = Path(file_path).suffix.lower()
    parsers = {
        ".csv": parse_inbody_from_csv,
        ".xlsx": parse_inbody_from_excel,
        ".xls": parse_inbody_from_excel,
        ".pdf": parse_inbody_from_pdf,
        ".docx": parse_inbody_from_docx,
        ".png": parse_inbody_with_ollama,
        ".jpg": parse_inbody_with_ollama,
        ".jpeg": parse_inbody_with_ollama,
    }

    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"지원하지 않는 파일 형식: {ext} (지원: {list(parsers.keys())})")

    return parser(file_path, **kwargs)


# ═══════════════════════════════════════
# CSV/Excel 행 매핑 유틸리티
# ═══════════════════════════════════════

# CSV/Excel 컬럼명 → 표준 필드 매핑 (괄호 포함 변형 모두 지원)
_CSV_COLUMN_MAP = {
    # 날짜
    "날짜": "test_date", "검사일": "test_date", "측정일": "test_date", "date": "test_date",
    "검사일시": "test_date", "측정일시": "test_date",
    # 키
    "키": "height_cm", "신장": "height_cm", "height": "height_cm",
    "신장(cm)": "height_cm", "키(cm)": "height_cm",
    # 나이
    "나이": "age", "연령": "age", "age": "age",
    # 성별
    "성별": "gender", "gender": "gender",
    # 체중
    "체중": "weight_kg", "weight": "weight_kg",
    "체중(kg)": "weight_kg",
    # 골격근량
    "골격근량": "skeletal_muscle_mass_kg", "smm": "skeletal_muscle_mass_kg",
    "골격근량(kg)": "skeletal_muscle_mass_kg",
    # 체지방량
    "체지방량": "body_fat_mass_kg", "체지방량(kg)": "body_fat_mass_kg", "체지방(kg)": "body_fat_mass_kg",
    # 체지방률
    "체지방률": "body_fat_percent", "bfp": "body_fat_percent",
    "체지방률(%)": "body_fat_percent",
    # BMI
    "bmi": "bmi", "BMI": "bmi", "BMI(kg/m²)": "bmi", "BMI(kg/m2)": "bmi",
    # 기초대사량
    "기초대사량": "bmr_kcal", "bmr": "bmr_kcal",
    "기초대사량(kcal)": "bmr_kcal",
    # 체수분
    "체수분": "body_water_L", "체수분(L)": "body_water_L",
    # 단백질
    "단백질": "protein_kg", "단백질(kg)": "protein_kg",
    # 무기질
    "무기질": "minerals_kg", "무기질(kg)": "minerals_kg",
    # 인바디점수
    "인바디점수": "inbody_score", "점수": "inbody_score", "InBody점수": "inbody_score",
    # 내장지방
    "내장지방": "vfa_cm2", "VFA": "vfa_cm2",
    "내장지방단면적": "vfa_cm2", "내장지방단면적(cm²)": "vfa_cm2", "내장지방단면적(cm2)": "vfa_cm2",
    "VFA(cm²)": "vfa_cm2", "VFA(cm2)": "vfa_cm2",
    # 복부지방률
    "복부지방률": "waist_hip_ratio", "WHR": "waist_hip_ratio",
    # SMI
    "SMI": "smi", "smi": "smi", "SMI(kg/m²)": "smi", "SMI(kg/m2)": "smi",
    "세포외수분비": "ecw_ratio",
}


def _safe_float(value):
    """값을 안전하게 float로 변환합니다."""
    if value is None:
        return None
    try:
        s = str(value).strip().replace(",", "")
        if s == "" or s.lower() == "nan" or s.lower() == "none" or s == "-":
            return None
        return float(s)
    except (ValueError, TypeError):
        return None


def _fuzzy_column_match(col_name: str) -> str:
    """매핑 테이블에 없는 컬럼명을 키워드로 자동 매칭합니다."""
    col = col_name.lower().replace(" ", "").replace("_", "")
    fuzzy_rules = [
        ("신장", "height_cm"), ("키", "height_cm"), ("height", "height_cm"),
        ("나이", "age"), ("연령", "age"), ("age", "age"),
        ("성별", "gender"),
        ("체중", "weight_kg"), ("weight", "weight_kg"),
        ("골격근", "skeletal_muscle_mass_kg"), ("smm", "skeletal_muscle_mass_kg"),
        ("체지방량", "body_fat_mass_kg"), ("bodyfatmass", "body_fat_mass_kg"),
        ("체지방률", "body_fat_percent"), ("체지방%", "body_fat_percent"), ("bfp", "body_fat_percent"), ("percentbodyfat", "body_fat_percent"),
        ("bmi", "bmi"),
        ("기초대사", "bmr_kcal"), ("bmr", "bmr_kcal"),
        ("체수분", "body_water_L"),
        ("단백질", "protein_kg"),
        ("무기질", "minerals_kg"),
        ("인바디점수", "inbody_score"), ("inbodyscore", "inbody_score"), ("점수", "inbody_score"),
        ("내장지방", "vfa_cm2"), ("vfa", "vfa_cm2"),
        ("복부지방", "waist_hip_ratio"), ("whr", "waist_hip_ratio"),
        ("smi", "smi"),
        ("검사일", "test_date"), ("날짜", "test_date"), ("측정일", "test_date"),
        ("세포외수분비", "ecw_ratio"), ("ecw", "ecw_ratio"),
    ]
    for keyword, field in fuzzy_rules:
        if keyword in col:
            return field
    return col_name


def _map_csv_row_to_inbody(row: dict) -> dict:
    """CSV/Excel 행을 인바디 표준 JSON으로 변환합니다."""
    mapped = {}

    # 숫자여야 하는 필드 목록
    numeric_fields = {
        "height_cm", "age", "weight_kg", "skeletal_muscle_mass_kg",
        "body_fat_mass_kg", "body_fat_percent", "bmi", "bmr_kcal",
        "vfa_cm2", "waist_hip_ratio", "smi", "ecw_ratio",
        "inbody_score", "body_water_L", "protein_kg", "minerals_kg",
        "exercise_frequency",
    }

    for col, value in row.items():
        col_clean = str(col).strip()
        # 1차: 정확 매핑 → 2차: 퍼지 매칭
        field_name = _CSV_COLUMN_MAP.get(col_clean)
        if field_name is None:
            field_name = _fuzzy_column_match(col_clean)

        if value is not None and str(value).strip() != "" and str(value) != "nan":
            if field_name in numeric_fields:
                mapped[field_name] = _safe_float(value)
            else:
                try:
                    mapped[field_name] = float(value) if isinstance(value, (int, float)) else value
                except (ValueError, TypeError):
                    mapped[field_name] = value

    # 표준 구조로 변환
    return {
        "basic_info": {
            "height_cm": mapped.get("height_cm"),
            "age": mapped.get("age"),
            "gender": mapped.get("gender"),
            "test_date": str(mapped.get("test_date", "")),
        },
        "body_composition": {
            "body_water_L": mapped.get("body_water_L"),
            "protein_kg": mapped.get("protein_kg"),
            "minerals_kg": mapped.get("minerals_kg"),
            "body_fat_mass_kg": mapped.get("body_fat_mass_kg"),
            "weight_kg": mapped.get("weight_kg"),
        },
        "muscle_fat_analysis": {
            "skeletal_muscle_mass_kg": mapped.get("skeletal_muscle_mass_kg"),
            "body_fat_mass_kg": mapped.get("body_fat_mass_kg"),
            "weight_kg": mapped.get("weight_kg"),
        },
        "obesity_analysis": {
            "bmi": mapped.get("bmi"),
            "body_fat_percent": mapped.get("body_fat_percent"),
        },
        "weight_control": {
            "ideal_weight_kg": None,
            "weight_adjustment_kg": None,
            "fat_adjustment_kg": None,
            "muscle_adjustment_kg": None,
        },
        "research_params": {
            "bmr_kcal": mapped.get("bmr_kcal"),
            "waist_hip_ratio": mapped.get("waist_hip_ratio"),
            "smi": mapped.get("smi"),
        },
        "visceral_fat": {"vfa_cm2": mapped.get("vfa_cm2")},
        "inbody_score": mapped.get("inbody_score"),
    }


def _build_history_from_records(records: list[dict]) -> dict:
    """여러 측정 기록에서 히스토리를 구성합니다."""
    history = {"dates": [], "weight_series": [], "smm_series": [], "bfp_series": []}
    for r in records:
        date = r.get("basic_info", {}).get("test_date", "")
        weight = r.get("body_composition", {}).get("weight_kg")
        smm = r.get("muscle_fat_analysis", {}).get("skeletal_muscle_mass_kg")
        bfp = r.get("obesity_analysis", {}).get("body_fat_percent")
        history["dates"].append(str(date))
        history["weight_series"].append(weight)
        history["smm_series"].append(smm)
        history["bfp_series"].append(bfp)
    return history


# ═══════════════════════════════════════
# 텍스트(PDF/Word)에서 정규식 파싱
# ═══════════════════════════════════════

_TEXT_PATTERNS = {
    "height_cm": r"(?:신장|키)\s*[:=]?\s*(\d{2,3}(?:\.\d+)?)\s*(?:cm)?",
    "age": r"(?:나이|연령)\s*[:=]?\s*(\d{1,3})",
    "weight_kg": r"체중\s*[:=]?\s*(\d{2,3}(?:\.\d+)?)\s*(?:kg)?",
    "skeletal_muscle_mass_kg": r"골격근량?\s*[:=]?\s*(\d{1,3}(?:\.\d+)?)",
    "body_fat_mass_kg": r"체지방(?:량)?\s*[:=]?\s*(\d{1,3}(?:\.\d+)?)\s*(?:kg)?",
    "bmi": r"BMI\s*[:=]?\s*(\d{1,2}(?:\.\d+)?)",
    "body_fat_percent": r"체지방률\s*[:=]?\s*(\d{1,2}(?:\.\d+)?)\s*%?",
    "bmr_kcal": r"기초대사량\s*[:=]?\s*(\d{3,4})\s*(?:kcal)?",
    "inbody_score": r"(?:인바디|InBody)\s*(?:점수|Score)\s*[:=]?\s*(\d{2,3})",
    "vfa_cm2": r"(?:내장지방|VFA)\s*[:=]?\s*(\d{1,3}(?:\.\d+)?)\s*(?:cm)?",
    "smi": r"SMI\s*[:=]?\s*(\d{1,2}(?:\.\d+)?)",
    "waist_hip_ratio": r"복부지방률?\s*[:=]?\s*(\d\.\d+)",
}


def _parse_inbody_from_text(text: str, method: str) -> InBodyData:
    """텍스트에서 정규식으로 인바디 데이터를 추출합니다."""
    parsed = {}
    for field_name, pattern in _TEXT_PATTERNS.items():
        match = re.search(pattern, text)
        if match:
            val = match.group(1)
            try:
                parsed[field_name] = float(val) if "." in val else int(val)
            except ValueError:
                parsed[field_name] = val

    result = {
        "basic_info": {
            "height_cm": parsed.get("height_cm"),
            "age": parsed.get("age"),
            "gender": "남성" if "남" in text[:200] else "여성" if "여" in text[:200] else None,
            "test_date": None,
        },
        "body_composition": {
            "body_fat_mass_kg": parsed.get("body_fat_mass_kg"),
            "weight_kg": parsed.get("weight_kg"),
        },
        "muscle_fat_analysis": {
            "skeletal_muscle_mass_kg": parsed.get("skeletal_muscle_mass_kg"),
            "body_fat_mass_kg": parsed.get("body_fat_mass_kg"),
            "weight_kg": parsed.get("weight_kg"),
        },
        "obesity_analysis": {
            "bmi": parsed.get("bmi"),
            "body_fat_percent": parsed.get("body_fat_percent"),
        },
        "weight_control": {},
        "research_params": {
            "bmr_kcal": parsed.get("bmr_kcal"),
            "waist_hip_ratio": parsed.get("waist_hip_ratio"),
            "smi": parsed.get("smi"),
        },
        "visceral_fat": {"vfa_cm2": parsed.get("vfa_cm2")},
        "inbody_score": parsed.get("inbody_score"),
    }

    return _build_inbody_data(result, method)


# ═══════════════════════════════════════
# 공통 빌더
# ═══════════════════════════════════════

def _build_inbody_data(data: dict, method: str) -> InBodyData:
    """딕셔너리를 InBodyData로 변환하고 교차 검증합니다."""
    total_fields = 0
    non_null_fields = 0
    for section_name in ["basic_info", "body_composition", "muscle_fat_analysis",
                         "obesity_analysis", "weight_control", "research_params"]:
        section = data.get(section_name, {})
        if isinstance(section, dict):
            for v in section.values():
                total_fields += 1
                if v is not None:
                    non_null_fields += 1

    confidence = round(non_null_fields / max(total_fields, 1), 2)
    validation = validate_inbody_data(data)

    return InBodyData(
        raw=data,
        basic_info=data.get("basic_info", {}),
        body_composition=data.get("body_composition", {}),
        muscle_fat_analysis=data.get("muscle_fat_analysis", {}),
        obesity_analysis=data.get("obesity_analysis", {}),
        weight_control=data.get("weight_control", {}),
        research_params=data.get("research_params", {}),
        segmental_lean=data.get("segmental_lean", {}),
        ecw_ratio=data.get("ecw_ratio", {}),
        body_history=data.get("body_history", {}),
        visceral_fat=data.get("visceral_fat", {}),
        inbody_score=data.get("inbody_score"),
        parse_method=method,
        parse_confidence=confidence,
        validation=validation,
    )
